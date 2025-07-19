
import streamlit as st
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

weights_map = {
    "VEG 1": ("25", "25"),
    "VEG 2": ("40", "40"),
    "VEG 3": ("60", "60"),
    "VEG 4": ("160", "220"),
    "PROTEIN 1": ("60", "70"),
    "GARNISH 1": ("1", "2"),
    "SAUCE 1": ("8", "8")
}

def extract_steps(text, meal_code):
    steps = []
    for line in text.strip().split("\n"):
        match = re.match(r"(\d+) - ([^:]+): \([^)]+\) (.+?) \u2192 Step (\d+)\) (P\d)", line.strip())
        if match:
            step_num, comp_type_raw, name, step_number, step_code = match.groups()
            component_type = comp_type_raw.strip().upper()
            placement = "Position " + step_code[-1]
            step_id = f"{step_num}-{step_code}"
            std_wt, lg_wt = weights_map.get(component_type, ("", ""))
            steps.append({
                "Component Type": component_type,
                "Meal Component Name": name.strip(),
                "Placement": placement,
                "Step": step_id,
                "Meal Code": meal_code,
                "Standard Size": std_wt,
                "Large Size": lg_wt
            })
    return steps

def add_step_page(doc, step):
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    for col in table.columns:
        col.width = Inches(1.2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = f"COMPONENT\n(TYPE)\n{step['Component Type']}"
    hdr_cells[1].text = f"PLACEMENT\n{step['Placement']}"
    hdr_cells[2].text = f"STEP (#)\n{step['Step']}"
    hdr_cells[3].text = f"MEAL CODE\n{step['Meal Code']}"

    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.runs[0]
            run.font.size = Pt(12)
            run.bold = True

    doc.add_paragraph("\nMEAL COMPONENT NAME", style='Heading 2').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_para = doc.add_paragraph(step['Meal Component Name'])
    name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_para.runs[0].font.size = Pt(14)
    name_para.runs[0].bold = True

    doc.add_paragraph("\nSTANDARD (SIZE)", style='Heading 2').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    std_para = doc.add_paragraph(step['Standard Size'])
    std_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    std_para.runs[0].font.size = Pt(28)
    std_para.runs[0].bold = True

    doc.add_paragraph("\nLARGE (SIZE)", style='Heading 2').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    lg_para = doc.add_paragraph(step['Large Size'])
    lg_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    lg_para.runs[0].font.size = Pt(28)
    lg_para.runs[0].bold = True

    doc.add_page_break()

st.title("Plating Step PDF Generator")
st.write("Paste your full plating instructions. One step per line.")

meal_code = st.text_input("Meal Code", value="A")
num_steps = st.number_input("How many steps are you entering?", min_value=1, max_value=100, value=15)

plating_text = st.text_area("Plating Steps", height=300)

if st.button("Generate PDF"):
    steps = extract_steps(plating_text, meal_code)

    if not steps:
        st.error("No valid steps found. Please check your formatting.")
    else:
        doc = Document()
        for step in steps[:num_steps]:
            add_step_page(doc, step)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button("📄 Download DOCX", data=buffer, file_name="plating_steps.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
